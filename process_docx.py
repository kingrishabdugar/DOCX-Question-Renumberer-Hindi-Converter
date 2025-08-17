import re
from docx import Document

# Hindi digit mapping
HINDI_DIGITS = str.maketrans('0123456789', '०१२३४५६७८९')

def renumber_and_convert_docx(input_path, output_path):
    doc = Document(input_path)
    question_counter = 1
    
    # Pass 1: Find and renumber questions (in order they appear)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            # Find patterns like "प्र: 98" or "प्र:98" (with optional space)
            matches = re.findall(r'(प्र:\s*\d+)', text)
            if matches:
                for match in matches:
                    # Extract the number part after "प्र:" (ignore old number, just renumber)
                    new_label = f"प्र: {question_counter}"
                    text = text.replace(match, new_label)
                    question_counter += 1
                run.text = text
    
    # Pass 2: Convert all English digits to Hindi in the entire document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.translate(HINDI_DIGITS)
    
    doc.save(output_path)

# Example usage (uncomment to test locally)
# renumber_and_convert_docx('input.docx', 'output.docx')
