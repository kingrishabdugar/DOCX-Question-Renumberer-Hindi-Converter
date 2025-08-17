import streamlit as st
import re
from docx import Document
from io import BytesIO

# Hindi digit mapping
HINDI_DIGITS = str.maketrans('0123456789', '०१२३४५६७८९')

def process_docx(file):
    doc = Document(file)
    question_counter = 1
    
    # Pass 1: Renumber questions
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            matches = re.findall(r'(प्र:\s*\d+)', text)
            if matches:
                for match in matches:
                    new_label = f"प्र: {question_counter}"
                    text = text.replace(match, new_label)
                    question_counter += 1
                run.text = text
    
    # Pass 2: Convert all English digits to Hindi
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.translate(HINDI_DIGITS)
    
    # Save to BytesIO for download
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("DOCX Question Renumberer & Hindi Converter")

uploaded_file = st.file_uploader("Upload your .docx file", type="docx")

if uploaded_file:
    if st.button("Process and Download"):
        with st.spinner("Processing..."):
            processed_file = process_docx(uploaded_file)
        st.download_button(
            label="Download Modified DOCX",
            data=processed_file,
            file_name="modified_output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
